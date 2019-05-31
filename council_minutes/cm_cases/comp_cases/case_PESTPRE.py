from docx import Document
from ...models import Request
from docx.enum.text import WD_ALIGN_PARAGRAPH

class PESTPRE():

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO(request, docx):
        para = docx.add_paragraph()
        para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        para.add_run('El Consejo de Facultad ')
        if request.approval_status == 'AP':
            PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_AP(request, docx, para)
        else:
            PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_NP(request, docx, para)
        

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_AP(request, docx, paragraph):
        paragraph.add_run('APRUEBA').font.bold = True
        paragraph.add_run(' inscribir la siguiente asignatura en el periódo académico ')
        paragraph.add_run(request.academic_period + ', a desarrollar en la empresa ')
        paragraph.add_run(request.detail_cm['institution'] + ', a cargo del ' + request.detail_cm['charge_un'])
        paragraph.add_run(' ' + request.detail_cm['person_un'] + ' por parte de la Universidad ')
        paragraph.add_run('Nacional de Colombia y')
        if request.detail_cm['title_ins'] == 'Sr.':
            paragraph.add_run(' el ')
        else:
            paragraph.add_run(' la ')
        paragraph.add_run(request.detail_cm['title_ins'] + ' ' + request.detail_cm['person_ins'])
        paragraph.add_run(' por parte de la entidad.')
        paragraph.add_run(' (Artículo 15 Acuerdo 008 de 2008 del Consejo Superior Universitario).')
        PESTPRE.case_PRACTICA_ESTUDIANTIL_PREGRADO_TABLE(request, docx)

    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_NP(request, docx, paragraph):
        paragraph.add_run('NO APRUEBA').font.bold = True
        paragraph.add_run(' inscribir la asignatura Práctica estudianti debido a que  {}'.format(request.justification))
        paragraph.add_run(' (Acuerdo 102 de 2013 del Consejo Superior Universitario).')
        
    @staticmethod
    def case_PRACTICA_ESTUDIANTIL_PREGRADO_TABLE(request, docx):
        table = docx.add_table(rows=len(request.detail_cm['subjects'])+1, cols=5, style='Table Grid')
        cellp = table.cell(0, 0).paragraphs[0]
        cellp.add_run('Código SIA').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 1).paragraphs[0]
        cellp.add_run('Nombre Asignatura').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 2).paragraphs[0]
        cellp.add_run('Grupo').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 3).paragraphs[0]
        cellp.add_run('T').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        cellp = table.cell(0, 4).paragraphs[0]
        cellp.add_run('C').font.bold = True
        cellp.alignment = WD_ALIGN_PARAGRAPH.CENTER

        index = 0
        for subject in request.detail_cm['subjects']:
            table.cell(index+1, 0).paragraphs[0].add_run(subject['code'])
            table.cell(index+1, 1).paragraphs[0].add_run(subject['subject'])
            table.cell(index+1, 4).paragraphs[0].add_run(subject['group'])
            table.cell(index+1, 3).paragraphs[0].add_run(subject['tipology'])
            table.cell(index+1, 2).paragraphs[0].add_run(subject['credits'])
            index = index + 1