from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, EmbeddedDocumentListField, BooleanField, DateField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph, num_to_month


class EMSP(Request):

    full_name = 'Exención por mejor saber pro'

    regulation_list = ['002|2011|CFA']  # List of regulations

    percentaje = FloatField(display='Porcentaje de exención de matrícula')
    is_renovation = BooleanField(display='Es renovación de beca')
    is_best = BooleanField(
        display='La beca es para el mejor puntaje, sino, es dentro de los primeros 10')
    date_presentation = DateField(display='Fecha de presentación del examen')

    str_cm = [
        'BECA EXCENCIÓN DE DERECHOS ACADÉMICOS del programa {} ({}) por obtener un excelente resu' +
        'ltado en el exámen de estado SABER-PRO.'
    ]

    str_pcm_modifiers = [
        'la renovación de ',
        'al mejor puntaje',
        'a los 10 mejores puntajes',
    ]

    str_pcm = [
        'la exención del {}% del pago de los derechos académicos y renovación de matrícula {} de' +
        'l Examen de Estado de la Calidad de la Educación Superior (SABER PRO) {} - {} para el pe' +
        'riodo académico {}. (Literal b, Artículo 16 del {}).'
    ]

    def cm(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        paragraph.add_run(self.str_cm[0].format(
            # pylint: disable=no-member
            self.get_academic_program_display(), self.academic_program))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_answer + ': ').bold = True
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        paragraph.add_run(
            self.str_pcm_modifiers[0] if self.is_renovation else '')
        paragraph.add_run(
            self.str_pcm[0].format(
                # pylint: disable=no-member
                self.percentaje,
                self.str_pcm_modifiers[1] if self.is_best else self.str_pcm_modifiers[2],
                self.date_presentation.year,
                self.get_academic_program_display(),
                self.academic_period,
                self.regulations['002|2011|CFA'][0]
            ))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)
