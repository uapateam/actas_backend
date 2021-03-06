from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from mongoengine import StringField, IntField, FloatField, EmbeddedDocumentListField, BooleanField, DateField
from ..models import Request, Subject
from .case_utils import table_subjects, add_analysis_paragraph, num_to_month


class CAIM(Request):

    full_name = 'Cancelación de asignaturas con carga inferior a la mínima'

    percentaje = FloatField(display='Porcentaje de avance de carrrera')
    enrollments = IntField(display='Número de matrículas')
    gpa = FloatField(display='P.A.P.A.')
    available_credits = IntField(display='Créditos disponibles')
    remaining_credits = IntField(display='Créditos restantes')
    subjects = EmbeddedDocumentListField(
        Subject, required=True, display='Asignaturas')

    regulation_list = ['008|2008|CSU']  # List of regulations

    str_cm = [
        'cursar el periodo académico {} con un número de créditos inferior al mínimo exigido, ',
        'cancelar la(s) siguiente(s) asignatura(s) inscrita(s) del periodo {}.',
        'debido a que {}realiza debidamente la solicitud.',
        '(Artículo 10 del {}).',
        '(Artículo 15 del {}).'
    ]

    str_pcm = [
        'SIA: Porcentaje de avance en el plan: {}%. Número de matrículas: {}. P.A.P.A.: {}.',
        'SIA: Créditos disponibles: {}. ',
        'SIA: Al aprobar la cancelación de la(s) asignatura(s) solicitada(s) el estudiante ' +
        'quedaría con {} créditos inscritos.'
    ]

    def cm(self, docx):
        if self.is_affirmative_response_approval_status():
            self.cm_ap(docx)
        else:
            self.cm_na(docx)

    def cm_ap(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ':')
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        self.cm_answer(paragraph)
        paragraph.add_run(self.str_cm[2].format('') + ' ')
        paragraph.add_run(self.str_cm[3].format(
            Request.regulations[self.regulation_list[0]][0]))
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[1].format(
            self.academic_period) + ' ')
        paragraph.add_run(self.str_cm[4].format(
            Request.regulations[self.regulation_list[0]][0]))
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def cm_na(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_council_header + ' ')
        self.cm_answer(paragraph)
        paragraph.add_run(self.str_cm[2].format('no ') + ' ')
        paragraph.add_run(self.str_cm[3].format(
            Request.regulations[self.regulation_list[0]][0]))

    def cm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_approval_status_display().upper() + ' ').font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0].format(self.academic_period))

    def pcm(self, docx):
        self.pcm_analysis(docx)
        if self.is_affirmative_response_advisor_response():
            self.pcm_ap(docx)
        else:
            self.pcm_na(docx)

    def pcm_ap(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ':')
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        self.pcm_answer(paragraph)
        paragraph.add_run(self.str_cm[2].format('') + ' ')
        paragraph.add_run(self.str_cm[3].format(
            Request.regulations[self.regulation_list[0]][0]))
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.style = 'List Bullet'
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper()).font.bold = True
        paragraph.add_run(' ' + self.str_cm[1].format(
            self.academic_period) + ' ')
        paragraph.add_run(self.str_cm[4].format(
            Request.regulations[self.regulation_list[0]][0]))
        table_subjects(docx, Subject.subjects_to_array(self.subjects))

    def pcm_na(self, docx):
        paragraph = docx.add_paragraph()
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        paragraph.paragraph_format.space_after = Pt(0)
        paragraph.add_run(self.str_comittee_header + ' ')
        self.pcm_answer(paragraph)
        paragraph.add_run(self.str_cm[2].format('no ') + ' ')
        paragraph.add_run(self.str_cm[3].format(
            Request.regulations[self.regulation_list[0]][0]))

    def pcm_answer(self, paragraph):
        paragraph.add_run(
            # pylint: disable=no-member
            self.get_advisor_response_display().upper() + ' ').font.bold = True
        # pylint: disable=no-member
        paragraph.add_run(self.str_cm[0].format(self.academic_period))

    def pcm_analysis(self, docx):
        analysis_list = []
        analysis_list += [self.str_pcm[0].format(
            self.percentaje,
            self.enrollments,
            self.gpa
        )]
        analysis_list += [self.str_pcm[1].format(
            self.available_credits
        )]
        analysis_list += [self.str_pcm[2].format(
            self.remaining_credits
        )]
        analysis_list += self.extra_analysis
        add_analysis_paragraph(docx, analysis_list)
